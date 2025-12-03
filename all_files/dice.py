from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_exam_docx():
    doc = Document()
    
    # --- Style Configurations ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # --- Header Section ---
    header_paragraphs = [
        "MERU UNIVERSITY OF SCIENCE AND TECHNOLOGY",
        "P.O. Box 972-60200 - Meru-Kenya",
        "Tel: +254(0) 799 529 958, +254(0) 799 529 959, +254 (0) 712 524 293",
        "Website: info@must.ac.ke Email: info@must.ac.ke",
        "",
        "University Examinations 2018/2019",
        "FIRST YEAR FIRST SEMESTER EXAMINATION FOR THE DEGREE OF BACHELOR OF SCIENCE IN MATHEMATICS AND COMPUTER SCIENCE, STATISTICS, ACTUARIAL SCIENCE, COMPUTER SCIENCE, COMPUTER TECHNOLOGY, INFORMATION TECHNOLOGY, INFORMATION SCIENCE, COMPUTER SCIENCE AND FORENSICS, HEALTH RECORDS AND INFORMATION MANAGEMENT, EDUCATION SCIENCE, ARTS EDUCATION, BIOLOGICAL SCIENCES, PHYSICAL SCIENCES, HEALTH SYSTEMS MANAGEMENT, PUBLIC HEALTH AND BUSINESS INFORMATION TECHNOLOGY.",
        "",
        "SMA 3110: MATHEMATICS FOR SCIENCE",
        "DATE: SEPTEMBER 2019                                                               TIME: 2 HOURS",
        "INSTRUCTIONS: Answer question one and any other two questions."
    ]

    for text in header_paragraphs:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if "MERU UNIVERSITY" in text or "SMA 3110" in text:
            p.runs[0].bold = True

    doc.add_paragraph("-" * 90)

    # --- Question 1 ---
    doc.add_paragraph("QUESTION ONE (20 MARKS)").runs[0].bold = True
    
    doc.add_paragraph("a) Briefly explain the following: (6 Marks)")
    doc.add_paragraph("    (i) Quadratic function.")
    doc.add_paragraph("    (ii) Permutation of objects.")
    doc.add_paragraph("    (iii) Conditional probability.")
    
    doc.add_paragraph("b) Solve x(x-3) < 10. (4 Marks)")
    
    p = doc.add_paragraph("c) Find the value of x such that: ")
    p.add_run(r"$(\frac{4}{5})^{5x-4}=(\frac{5}{4})^{4x-5}$").italic = True
    p.add_run(" (3 Marks)")

    doc.add_paragraph("d) In how many ways can a committee of 7 be selected from 10 men and 7 women if there must be a majority of women serving? (4 Marks)")

    p = doc.add_paragraph("e) Show that the sum of n terms of the series ")
    p.add_run(r"$log~a+log~2a+log~4a+log~8a+...$")
    p.add_run(" is ")
    # Preserving the complex nested power structure here strictly
    p.add_run(r"$log[a^{n}2^{n^{\frac{n-1}{2})}}]$").italic = True
    p.add_run(". (4 Marks)")

    p = doc.add_paragraph("f) If ")
    p.add_run(r"$sin~A=\frac{3}{4}$")
    p.add_run(r" and $90^{\circ}\le A\le180^{\circ}$")
    p.add_run(" find CotA, without using tables or calculator. (3 Marks)")

    doc.add_paragraph("g) Calculate the mean absolute deviation for 2, 11, 3, 7, 7. (3 Marks)")

    doc.add_paragraph("h) From past records, Wanyama has 75% success rate in taking penalty shots. In his next 6 shots, what is the probability that he succeeds in exactly 4 shots? (3 Marks)")

    # --- Question 2 ---
    doc.add_paragraph("\nQUESTION TWO (20 MARKS)").runs[0].bold = True
    
    p = doc.add_paragraph("a) Find the maximum value of the function ")
    p.add_run(r"$y=15+6x-3x^{2}$").italic = True
    p.add_run(" by writing it in the form ")
    p.add_run(r"$y=a(x-p)^{2}+q$").italic = True
    p.add_run(". (4 Marks)")

    p = doc.add_paragraph("b) Determine the values of n for which the equation ")
    p.add_run(r"$x^{2}+2(n+1)x+2(n+5)=0$").italic = True
    p.add_run(" has real roots. (3 Marks)")

    p = doc.add_paragraph("c) Solve ")
    p.add_run(r"$2x^{2}-5x-6=0$").italic = True
    p.add_run(" using completing square method. (3 Marks)")

    p = doc.add_paragraph("d) Simplify ")
    p.add_run(r"$\frac{2\sqrt{3}+4\sqrt{7}}{3\sqrt{5}-2\sqrt{2}}$").italic = True
    p.add_run(" by rationalizing the denominator. (3 Marks)")

    p = doc.add_paragraph("e) Find x if ")
    p.add_run(r"$log_{3}x-\frac{1}{2}=\frac{3}{2}log_{x}3$").italic = True
    p.add_run(". (4 Marks)")

    p = doc.add_paragraph("f) Show that ")
    p.add_run(r"$log_{16}x=\frac{1}{4}log_{2}x$").italic = True
    p.add_run(". (3 Marks)")

    # --- Question 3 ---
    doc.add_paragraph("\nQUESTION THREE (20 MARKS)").runs[0].bold = True

    doc.add_paragraph("a) In how many ways can one arrange the word STATISTICIAN in a row? (4 Marks)")

    doc.add_paragraph("b) There are 8 persons, including a married couple Mr. and Mrs. Sharma, from whom a committee of 4 has to be chosen. In how many ways can the committee be chosen: (3 Marks each)")
    doc.add_paragraph("    (i) If both Mr. and Mrs. Sharma are excluded.")
    doc.add_paragraph("    (ii) If Mrs. Sharma is included but Mr. Sharma is excluded.")
    doc.add_paragraph("    (iii) If both Mr. and Mrs. Sharma are included.")

    doc.add_paragraph("c) In an Arithmetic progression, the fourth term is 13 and the seventh term is 22. Determine:")
    p = doc.add_paragraph("    (i) The value of n if the ")
    p.add_run(r"$n^{th}$").italic = True
    p.add_run(" term is 100. (4 Marks)")
    doc.add_paragraph("    (ii) The value of m if the sum of first m terms of the series is 175. (3 Marks)")

    # --- Question 4 ---
    doc.add_paragraph("\nQUESTION FOUR (20 MARKS)").runs[0].bold = True

    p = doc.add_paragraph("a) Prove the following trigonometric identity: ")
    p.add_run(r"$Sin^{2}\theta~cot~\theta~sec~\theta=sin~\theta$").italic = True
    p.add_run(". (3 Marks)")

    p = doc.add_paragraph("b) Without using a calculator, simplify completely by giving your answer in the form ")
    p.add_run(r"$a+b\sqrt{z}$").italic = True
    doc.add_paragraph("    i) $cos~750^{\circ}$")
    doc.add_paragraph("    ii) $tan~405^{\circ}+tan~510^{\circ}$ (4 Marks)")

    p = doc.add_paragraph("c) Given that ")
    p.add_run(r"$cos~2\theta=1-sin^{2}\theta$").italic = True
    p.add_run(", solve the following trigonometric equation: ")
    p.add_run(r"$3~cos~2\vartheta+sin~\theta=1$").italic = True
    p.add_run(r" for $0^{\circ}\le\theta\le360$").italic = True
    p.add_run(". (5 Marks)")

    p = doc.add_paragraph("d) A triangle ABC has the following measures; AB=3.8cm, BC=5.5cm and AC=4.5cm. Determine the largest angle of the triangle ABC. (5 Marks)")

    p = doc.add_paragraph("e) Show that ")
    p.add_run(r"$2~sin~75~cos~45=\frac{1}{2}(\sqrt{3}+1)$").italic = True
    p.add_run(". (3 Marks)")

    # --- Question 5 ---
    doc.add_paragraph("\nQUESTION FIVE (20 MARKS)").runs[0].bold = True

    doc.add_paragraph("a) The mass of 100 students are shown in the following frequency distribution:")
    
    # Table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Mass (kg)'
    hdr_cells[1].text = 'Frequency'
    
    data = [
        ("50-59.99", "3"),
        ("60-69.99", "12"),
        ("70-79.99", "13"),
        ("80-89.99", "25"),
        ("90-99.99", "27"),
        ("100-109.99", "10"),
        ("110-119.99", "10")
    ]
    
    for mass, freq in data:
        row_cells = table.add_row().cells
        row_cells[0].text = mass
        row_cells[1].text = freq

    doc.add_paragraph("For this frequency distribution:")
    doc.add_paragraph("    (i) State the size of each class interval. (2 Marks)")
    doc.add_paragraph("    (ii) Calculate the mean and standard deviation. (6 Marks)")
    doc.add_paragraph("    (iii) Calculate the minimum acceptable mass if the heaviest 10% of the students are acceptable for special weight training. (4 Marks)")

    p = doc.add_paragraph("b) For two events A and B, ")
    p.add_run(r"$P(\sim A)=\frac{3}{10}, P(B)=\frac{13}{20}$").italic = True
    p.add_run(" and ")
    p.add_run(r"$P(A\cap B)=\frac{9}{20}$").italic = True
    p.add_run(". Find:")
    
    doc.add_paragraph("    (i) $P(A\\cup B)$ (2 Marks)")
    doc.add_paragraph("    (ii) $P(\\sim A\\cap B)$ (3 Marks)")
    doc.add_paragraph("    (iii) $P(A\\cup\\sim B)$ (3 Marks)")

    # Save
    doc.save('SMA_3110_Exam_Formatted.docx')

create_exam_docx()